"""Document parser supporting PDF, Word, and image files."""

import base64
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Any


@dataclass
class ParsedDocument:
    """Parsed document content."""

    file_path: str
    file_type: str  # pdf, docx, image, text
    text_content: str = ""
    images: List[dict] = field(default_factory=list)  # [{data: base64, type: mime}]
    pages: int = 0
    metadata: dict = field(default_factory=dict)

    def has_images(self) -> bool:
        return len(self.images) > 0

    def get_full_content(self) -> str:
        """Get text content with image placeholders."""
        if not self.images:
            return self.text_content
        return f"{self.text_content}\n\n[包含 {len(self.images)} 张图片]"


class DocumentParser:
    """Parse various document formats."""

    SUPPORTED_EXTENSIONS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".png": "image",
        ".jpg": "image",
        ".jpeg": "image",
        ".gif": "image",
        ".webp": "image",
        ".txt": "text",
        ".md": "text",
    }

    def __init__(self):
        pass

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse a document file.

        Args:
            file_path: Path to the document file

        Returns:
            ParsedDocument with extracted content
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        file_type = self.SUPPORTED_EXTENSIONS[ext]

        if file_type == "pdf":
            return self._parse_pdf(file_path)
        elif file_type == "docx":
            return self._parse_docx(file_path)
        elif file_type == "image":
            return self._parse_image(file_path)
        elif file_type == "text":
            return self._parse_text(file_path)
        else:
            raise ValueError(f"Unknown file type: {file_type}")

    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """Parse PDF file."""
        doc = ParsedDocument(
            file_path=file_path,
            file_type="pdf",
        )

        # Try PyMuPDF (fitz) first
        try:
            import fitz  # PyMuPDF

            pdf = fitz.open(file_path)
            doc.pages = len(pdf)
            doc.metadata = dict(pdf.metadata) if pdf.metadata else {}

            text_parts = []
            for page_num, page in enumerate(pdf):
                # Extract text
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"--- 第 {page_num + 1} 页 ---\n{text}")

                # Extract images
                for img_index, img in enumerate(page.get_images()):
                    try:
                        xref = img[0]
                        base_image = pdf.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        mime_type = f"image/{image_ext}"

                        doc.images.append({
                            "data": base64.b64encode(image_bytes).decode("utf-8"),
                            "type": mime_type,
                            "page": page_num + 1,
                            "index": img_index,
                        })
                    except Exception:
                        pass

            doc.text_content = "\n\n".join(text_parts)
            pdf.close()
            return doc

        except ImportError:
            pass

        # Fallback to pdfplumber
        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                doc.pages = len(pdf.pages)

                text_parts = []
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"--- 第 {page_num + 1} 页 ---\n{text}")

                doc.text_content = "\n\n".join(text_parts)
                return doc

        except ImportError:
            pass

        # Fallback to PyPDF2
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            doc.pages = len(reader.pages)

            text_parts = []
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(f"--- 第 {page_num + 1} 页 ---\n{text}")

            doc.text_content = "\n\n".join(text_parts)
            return doc

        except ImportError:
            raise ImportError(
                "No PDF parser available. Install one of: "
                "PyMuPDF (pip install pymupdf), "
                "pdfplumber (pip install pdfplumber), "
                "or PyPDF2 (pip install pypdf2)"
            )

    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse Word document."""
        doc = ParsedDocument(
            file_path=file_path,
            file_type="docx",
        )

        try:
            from docx import Document
            from docx.opc.constants import RELATIONSHIP_TYPE as RT

            word_doc = Document(file_path)

            # Extract text from paragraphs
            text_parts = []
            for para in word_doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in word_doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n".join(table_text))

            doc.text_content = "\n\n".join(text_parts)
            doc.pages = len(word_doc.paragraphs) // 30 + 1  # Rough estimate

            # Extract images
            for rel in word_doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        content_type = image_part.content_type

                        doc.images.append({
                            "data": base64.b64encode(image_bytes).decode("utf-8"),
                            "type": content_type,
                        })
                    except Exception:
                        pass

            return doc

        except ImportError:
            raise ImportError(
                "python-docx not installed. Run: pip install python-docx"
            )

    def _parse_image(self, file_path: str) -> ParsedDocument:
        """Parse image file."""
        doc = ParsedDocument(
            file_path=file_path,
            file_type="image",
            pages=1,
        )

        # Determine MIME type
        ext = Path(file_path).suffix.lower()
        mime_map = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".gif": "image/gif",
            ".webp": "image/webp",
        }
        mime_type = mime_map.get(ext, "image/png")

        # Read and encode image
        with open(file_path, "rb") as f:
            image_bytes = f.read()

        doc.images.append({
            "data": base64.b64encode(image_bytes).decode("utf-8"),
            "type": mime_type,
        })

        doc.text_content = f"[图片文件: {os.path.basename(file_path)}]"
        return doc

    def _parse_text(self, file_path: str) -> ParsedDocument:
        """Parse text file."""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        return ParsedDocument(
            file_path=file_path,
            file_type="text",
            text_content=content,
            pages=1,
        )
